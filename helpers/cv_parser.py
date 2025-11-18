import pdfplumber
import docx
import os

def extract_text(file_input):
    """
    Extract text from PDF, DOCX, or TXT.
    Works for both file paths and uploaded Streamlit file objects.
    """

    # If it's a Streamlit UploadedFile object
    if hasattr(file_input, "name"):
        filename = file_input.name
        ext = filename.split(".")[-1].lower()

        if ext == "pdf":
            return extract_text_from_pdf_streamlit(file_input)

        elif ext == "docx":
            return extract_text_from_docx_streamlit(file_input)

        elif ext == "txt":
            return file_input.read().decode("utf-8")

    # If it's a local file path
    else:
        ext = file_input.split(".")[-1].lower()

        if ext == "pdf":
            return extract_text_from_pdf(file_input)

        elif ext == "docx":
            return extract_text_from_docx(file_input)

        elif ext == "txt":
            return open(file_input, "r", encoding="utf-8").read()

    return ""


# ---------------- PDF Extract (Local File) ----------------

def extract_text_from_pdf(file_path):
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except:
        pass
    return text


# ---------------- PDF Extract (Streamlit Upload) ----------------

def extract_text_from_pdf_streamlit(uploaded_file):
    text = ""
    try:
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except:
        pass
    return text


# ---------------- DOCX Extract (Local File) ----------------

def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    except:
        return ""


# ---------------- DOCX Extract (Streamlit Upload) ----------------

def extract_text_from_docx_streamlit(uploaded_file):
    try:
        doc = docx.Document(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs])
    except:
        return ""
